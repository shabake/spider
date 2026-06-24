"""
文档格式转换工具

支持：
- Word → PDF
- PDF → Word

自动检测可用的转换后端（LibreOffice > python-docx+PyMuPDF）
"""

import os
import asyncio
import shutil
import tempfile


def _check_libreoffice() -> bool:
    """检查系统是否安装了 LibreOffice"""
    return shutil.which("libreoffice") is not None


def _check_python_libs() -> bool:
    """检查 Python 转换库是否安装"""
    try:
        import docx
        import fitz
        return True
    except ImportError:
        return False


AVAILABLE_BACKENDS = []

if _check_libreoffice():
    AVAILABLE_BACKENDS.append("libreoffice")
if _check_python_libs():
    AVAILABLE_BACKENDS.append("python")


async def docx_to_pdf(input_path: str, output_path: str = None) -> str:
    """
    将 Word 文档 (.docx) 转换为 PDF

    Args:
        input_path: 输入 .docx 文件路径
        output_path: 输出 .pdf 文件路径（可选，默认同目录同名）

    Returns:
        转换结果信息
    """
    input_path = os.path.expanduser(input_path)

    if not os.path.exists(input_path):
        return f"❌ 文件不存在: {input_path}"
    if not input_path.lower().endswith(".docx"):
        return f"❌ 不支持的文件格式: {input_path}，需要 .docx"

    if not output_path:
        output_path = os.path.splitext(input_path)[0] + ".pdf"
    output_path = os.path.expanduser(output_path)

    # 尝试 LibreOffice（最可靠，保留格式）
    if "libreoffice" in AVAILABLE_BACKENDS:
        return await _convert_via_libreoffice(input_path, output_path, "pdf")

    # 尝试 Python 库（基本转换，格式简单）
    if "python" in AVAILABLE_BACKENDS:
        return await _convert_via_python_docx_to_pdf(input_path, output_path)

    return (
        "❌ 没有可用的转换引擎\n\n"
        "请安装任一依赖：\n"
        "  # 方案 A（推荐，格式最好）:\n"
        "    brew install libreoffice\n\n"
        "  # 方案 B（轻量，格式简单）:\n"
        "    pip3 install python-docx PyMuPDF fpdf2"
    )


async def pdf_to_docx(input_path: str, output_path: str = None) -> str:
    """
    将 PDF 转换为 Word 文档 (.docx)

    Args:
        input_path: 输入 .pdf 文件路径
        output_path: 输出 .docx 文件路径（可选，默认同目录同名）

    Returns:
        转换结果信息
    """
    input_path = os.path.expanduser(input_path)

    if not os.path.exists(input_path):
        return f"❌ 文件不存在: {input_path}"
    if not input_path.lower().endswith(".pdf"):
        return f"❌ 不支持的文件格式: {input_path}，需要 .pdf"

    if not output_path:
        output_path = os.path.splitext(input_path)[0] + ".docx"
    output_path = os.path.expanduser(output_path)

    # 尝试 LibreOffice（最可靠）
    if "libreoffice" in AVAILABLE_BACKENDS:
        return await _convert_via_libreoffice(input_path, output_path, "docx")

    # 尝试 Python 库
    if "python" in AVAILABLE_BACKENDS:
        return await _convert_via_python_pdf_to_docx(input_path, output_path)

    return (
        "❌ 没有可用的转换引擎\n\n"
        "请安装任一依赖：\n"
        "  # 方案 A（推荐，格式最好）:\n"
        "    brew install libreoffice\n\n"
        "  # 方案 B（轻量，基本文本）:\n"
        "    pip3 install python-docx PyMuPDF pdf2docx"
    )


async def _convert_via_libreoffice(input_path: str, output_path: str, fmt: str) -> str:
    """使用 LibreOffice 进行转换（保留原始排版）"""
    output_dir = os.path.dirname(output_path) or "."
    cmd = [
        "libreoffice", "--headless", "--convert-to", fmt,
        "--outdir", output_dir, input_path
    ]
    proc = await asyncio.create_subprocess_exec(
        *cmd, stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE
    )
    stdout, stderr = await proc.communicate()

    if proc.returncode == 0:
        # LibreOffice 会用自己的文件名规则，检查是否需要重命名
        lo_output = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + "." + fmt)
        if lo_output != output_path and os.path.exists(lo_output):
            os.rename(lo_output, output_path)
        return f"✅ 转换成功 → {output_path}（使用 LibreOffice）"
    else:
        error = stderr.decode("utf-8", errors="replace")[:200]
        return f"❌ LibreOffice 转换失败: {error}"


async def _convert_via_python_docx_to_pdf(input_path: str, output_path: str) -> str:
    """使用 Python 库转换 docx → pdf（纯文本，格式较简单）"""
    try:
        from docx import Document
        from fpdf import FPDF
    except ImportError:
        return "❌ 需要安装: pip3 install python-docx fpdf2"

    try:
        doc = Document(input_path)

        pdf = FPDF()
        pdf.add_page()

        # macOS 系统中文字体
        font_paths = [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/AppleLiGothic.ttf",
        ]
        font_added = False
        for fp in font_paths:
            if os.path.exists(fp):
                pdf.add_font("CJK", "", fp, uni=True)
                font_added = True
                break

        if font_added:
            pdf.set_font("CJK", "", 12)
        else:
            pdf.set_font("Helvetica", "", 12)  # 仅英文

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 检查是否含中文
                has_cjk = any("一" <= c <= "鿿" for c in text)
                if has_cjk and not font_added:
                    pdf.set_font("Helvetica", "", 12)
                pdf.multi_cell(0, 8, text)
                pdf.ln(2)

        pdf.output(output_path)
        return f"✅ 转换成功 → {output_path}（使用 Python，基础格式）"
    except Exception as e:
        return f"❌ 转换失败: {e}"


async def _convert_via_python_pdf_to_docx(input_path: str, output_path: str) -> str:
    """使用 Python 库转换 pdf → docx"""
    try:
        import fitz  # PyMuPDF
        from docx import Document
    except ImportError:
        return "❌ 需要安装: pip3 install PyMuPDF python-docx"

    try:
        doc = Document()
        pdf = fitz.open(input_path)

        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text = page.get_text()

            if page_num > 0:
                doc.add_page_break()
            if text.strip():
                doc.add_paragraph(text.strip())

        doc.save(output_path)
        pdf.close()
        return f"✅ 转换成功 → {output_path}（使用 Python，文本提取）"
    except Exception as e:
        return f"❌ 转换失败: {e}"


# ── Tool Schema ─────────────────────────────────────────────

DOCX_TO_PDF_SCHEMA = {
    "type": "object",
    "properties": {
        "input_path": {
            "type": "string",
            "description": "输入的 Word 文件路径 (.docx)",
        },
        "output_path": {
            "type": "string",
            "description": "输出的 PDF 文件路径（可选，默认同目录同名）",
        },
    },
    "required": ["input_path"],
}

PDF_TO_DOCX_SCHEMA = {
    "type": "object",
    "properties": {
        "input_path": {
            "type": "string",
            "description": "输入的 PDF 文件路径",
        },
        "output_path": {
            "type": "string",
            "description": "输出的 Word 文件路径（可选，默认同目录同名）",
        },
    },
    "required": ["input_path"],
}
